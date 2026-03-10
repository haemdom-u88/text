import html
import io
import json
import os
from datetime import datetime

from flask import Blueprint, jsonify, request, send_file

import app_state
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

bp = Blueprint("graph", __name__)


@bp.route("/api/save_graph", methods=["POST"])
def save_graph():
    try:
        data = request.get_json() or {}
        nodes = data.get("nodes") or []
        edges = data.get("edges") or []
        if not nodes and not edges:
            return jsonify({"success": False, "error": "图谱为空"}), 400
        if not app_state.neo4j_store:
            return jsonify({"success": False, "error": "Neo4j 未配置，已跳过入库"}), 400
        app_state.neo4j_store.upsert_nodes_edges(nodes, edges)
        return jsonify({
            "success": True,
            "saved_nodes": len(nodes),
            "saved_edges": len(edges)
        })
    except Exception as exc:
        app_state.logger.error("Save graph failed: %s", exc, exc_info=True)
        return jsonify({"success": False, "error": str(exc)}), 500


@bp.route("/api/neo4j/subgraph", methods=["GET"])
def neo4j_subgraph():
    try:
        if not app_state.neo4j_store:
            return jsonify({"success": False, "error": "Neo4j 未配置"}), 400
        center = request.args.get("center")
        depth = int(request.args.get("depth", 1))
        limit_nodes = int(request.args.get("limit_nodes", 200))
        limit_edges = int(request.args.get("limit_edges", 800))
        include_props = str(request.args.get("include_props", "")).strip() in {"1", "true", "yes"}
        result = app_state.neo4j_store.fetch_subgraph(
            center=center,
            depth=depth,
            limit_nodes=limit_nodes,
            limit_edges=limit_edges,
            include_props=include_props
        )
        if result is None:
            return jsonify({"success": False, "error": "Neo4j 未配置"}), 400
        return jsonify({"success": True, "nodes": result.get("nodes", []), "edges": result.get("edges", [])})
    except Exception as exc:
        app_state.logger.error("Neo4j subgraph failed: %s", exc, exc_info=True)
        return jsonify({"success": False, "error": str(exc)}), 500


@bp.route("/api/export", methods=["POST"])
def export_data():
    try:
        data = request.get_json() or {}
        format_type = (data.get("format") or "").lower().strip()
        payload = data.get("data")
        if not format_type:
            return jsonify({"success": False, "error": "缺少导出格式"}), 400
        if format_type == "json":
            return jsonify({"success": True, "data": payload})

        os.makedirs("output", exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        text = json.dumps(payload, ensure_ascii=False, indent=2)

        if format_type == "pdf":
            filename = f"analysis_report_{timestamp}.pdf"
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = [Paragraph("Knowledge Graph Report", styles["Title"]), Spacer(1, 12)]
            for line in text.splitlines():
                safe_line = html.escape(line).replace(" ", "&nbsp;")
                story.append(Paragraph(safe_line, styles["Code"]))
            doc.build(story)
            buffer.seek(0)
            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype="application/pdf"
            )

        if format_type in {"word", "doc", "docx"}:
            filename = f"analysis_report_{timestamp}.docx"
            document = Document()
            document.add_heading("Knowledge Graph Report", level=1)
            document.add_paragraph(text)
            buffer = io.BytesIO()
            document.save(buffer)
            buffer.seek(0)
            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        return jsonify({"success": False, "error": f"不支持的格式: {format_type}"}), 400
    except Exception as exc:
        app_state.logger.error("Export failed: %s", exc, exc_info=True)
        return jsonify({"success": False, "error": f"导出失败: {exc}"}), 500
